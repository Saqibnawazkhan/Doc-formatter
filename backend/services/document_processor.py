import os
import re
import uuid
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from models.formatting_options import (
    FormattingOptions,
    TextAlignment,
    LineSpacing,
    PageSize,
    FontFamily,
    PageNumberPosition,
)


class DocumentProcessor:
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)

    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """Save uploaded file and return file_id"""
        file_id = str(uuid.uuid4())
        file_extension = os.path.splitext(filename)[1]
        file_path = os.path.join(self.upload_dir, f"{file_id}{file_extension}")

        with open(file_path, "wb") as f:
            f.write(file_content)

        return file_id

    def get_file_path(self, file_id: str) -> Optional[str]:
        """Get the full path for a file_id"""
        for ext in [".docx", ".doc"]:
            path = os.path.join(self.upload_dir, f"{file_id}{ext}")
            if os.path.exists(path):
                return path
        return None

    def format_document(self, file_id: str, options: FormattingOptions) -> str:
        """Apply formatting options to a document and return new file_id"""
        source_path = self.get_file_path(file_id)
        if not source_path:
            raise FileNotFoundError(f"File not found: {file_id}")

        doc = Document(source_path)

        # Apply formatting options
        if options.text:
            self._apply_text_formatting(doc, options.text)

        if options.paragraph:
            self._apply_paragraph_formatting(doc, options.paragraph)

        if options.page:
            self._apply_page_formatting(doc, options.page)

        if options.structure:
            self._apply_structure_formatting(doc, options.structure)

        if options.cleanup:
            self._apply_cleanup(doc, options.cleanup)

        # Save formatted document
        formatted_file_id = str(uuid.uuid4())
        formatted_path = os.path.join(self.upload_dir, f"{formatted_file_id}_formatted.docx")
        doc.save(formatted_path)

        return formatted_file_id

    def _apply_text_formatting(self, doc: Document, options):
        """Apply text formatting to all paragraphs"""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                try:
                    if options.font_family:
                        run.font.name = options.font_family.value
                        # Set font for East Asian characters safely
                        r = run._element
                        rPr = r.rPr
                        if rPr is not None:
                            rFonts = rPr.rFonts
                            if rFonts is not None:
                                rFonts.set(qn('w:eastAsia'), options.font_family.value)

                    if options.font_size:
                        run.font.size = Pt(options.font_size)

                    if options.font_color:
                        color = self._parse_color(options.font_color)
                        if color:
                            run.font.color.rgb = color

                    if options.bold is not None:
                        run.font.bold = options.bold

                    if options.italic is not None:
                        run.font.italic = options.italic

                    if options.underline is not None:
                        run.font.underline = options.underline
                except Exception:
                    continue

            if options.line_spacing:
                self._set_line_spacing(paragraph, options.line_spacing)

            if options.text_alignment:
                self._set_alignment(paragraph, options.text_alignment)

        # Also apply to tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                if options.font_family:
                                    run.font.name = options.font_family.value
                                if options.font_size:
                                    run.font.size = Pt(options.font_size)
                            except Exception:
                                continue

    def _apply_paragraph_formatting(self, doc: Document, options):
        """Apply paragraph formatting"""
        for paragraph in doc.paragraphs:
            pf = paragraph.paragraph_format

            if options.spacing_before is not None:
                pf.space_before = Pt(options.spacing_before)

            if options.spacing_after is not None:
                pf.space_after = Pt(options.spacing_after)

            if options.indent_left is not None:
                pf.left_indent = Inches(options.indent_left)

            if options.indent_right is not None:
                pf.right_indent = Inches(options.indent_right)

            if options.first_line_indent is not None:
                pf.first_line_indent = Inches(options.first_line_indent)

        if options.remove_extra_spaces:
            self._remove_extra_spaces(doc)

        if options.remove_blank_lines:
            self._remove_blank_lines(doc)

    def _apply_page_formatting(self, doc: Document, options):
        """Apply page formatting to all sections"""
        for section in doc.sections:
            if options.page_size:
                if options.page_size == PageSize.A4:
                    section.page_width = Inches(8.27)
                    section.page_height = Inches(11.69)
                elif options.page_size == PageSize.LETTER:
                    section.page_width = Inches(8.5)
                    section.page_height = Inches(11)

            if options.margin_top is not None:
                section.top_margin = Inches(options.margin_top)

            if options.margin_bottom is not None:
                section.bottom_margin = Inches(options.margin_bottom)

            if options.margin_left is not None:
                section.left_margin = Inches(options.margin_left)

            if options.margin_right is not None:
                section.right_margin = Inches(options.margin_right)

            if options.header_text:
                header = section.header
                header.is_linked_to_previous = False
                if header.paragraphs:
                    header.paragraphs[0].text = options.header_text
                else:
                    header.add_paragraph(options.header_text)

            if options.footer_text:
                footer = section.footer
                footer.is_linked_to_previous = False
                if footer.paragraphs:
                    footer.paragraphs[0].text = options.footer_text
                else:
                    footer.add_paragraph(options.footer_text)

            if options.page_numbers:
                self._add_page_numbers(section, options.page_number_position)

    def _apply_structure_formatting(self, doc: Document, options):
        """Apply document structure formatting"""
        if options.normalize_headings:
            self._normalize_headings(doc, options)

    def _apply_cleanup(self, doc: Document, options):
        """Apply cleanup and standardization"""
        if options.remove_inconsistent_fonts and options.normalize_formatting:
            default_font = "Calibri"
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    try:
                        run.font.name = default_font
                    except Exception:
                        continue

        if options.remove_extra_spaces or options.clean_copied_text:
            self._remove_extra_spaces(doc)

        if options.fix_alignment_issues:
            for paragraph in doc.paragraphs:
                if paragraph.paragraph_format.alignment is None:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _parse_color(self, color_str: str) -> Optional[RGBColor]:
        """Parse color string to RGBColor"""
        if color_str.startswith("#"):
            color_str = color_str[1:]

        if len(color_str) == 6:
            try:
                r = int(color_str[0:2], 16)
                g = int(color_str[2:4], 16)
                b = int(color_str[4:6], 16)
                return RGBColor(r, g, b)
            except ValueError:
                return None
        return None

    def _set_line_spacing(self, paragraph, spacing: LineSpacing):
        """Set line spacing for a paragraph"""
        pf = paragraph.paragraph_format
        spacing_map = {
            LineSpacing.SINGLE: 1.0,
            LineSpacing.ONE_POINT_FIFTEEN: 1.15,
            LineSpacing.ONE_POINT_FIVE: 1.5,
            LineSpacing.DOUBLE: 2.0,
        }
        if spacing in spacing_map:
            pf.line_spacing = spacing_map[spacing]

    def _set_alignment(self, paragraph, alignment: TextAlignment):
        """Set text alignment for a paragraph"""
        alignment_map = {
            TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
            TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
            TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
            TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if alignment in alignment_map:
            paragraph.paragraph_format.alignment = alignment_map[alignment]

    def _remove_extra_spaces(self, doc: Document):
        """Remove extra spaces from text"""
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                try:
                    run.text = re.sub(r' +', ' ', run.text)
                    run.text = run.text.strip()
                except Exception:
                    continue

    def _remove_blank_lines(self, doc: Document):
        """Remove consecutive blank paragraphs"""
        paragraphs_to_remove = []
        prev_blank = False

        for i, paragraph in enumerate(doc.paragraphs):
            is_blank = not paragraph.text.strip()
            if is_blank and prev_blank:
                paragraphs_to_remove.append(paragraph)
            prev_blank = is_blank

        for paragraph in paragraphs_to_remove:
            try:
                p = paragraph._element
                p.getparent().remove(p)
            except Exception:
                continue

    def _normalize_headings(self, doc: Document, options):
        """Normalize heading styles"""
        heading_sizes = {
            'Heading 1': options.h1_size or 24,
            'Heading 2': options.h2_size or 20,
            'Heading 3': options.h3_size or 16,
        }

        for paragraph in doc.paragraphs:
            try:
                if paragraph.style and paragraph.style.name in heading_sizes:
                    for run in paragraph.runs:
                        run.font.size = Pt(heading_sizes[paragraph.style.name])
                        if options.heading_font_family:
                            run.font.name = options.heading_font_family.value
                        run.font.bold = True
            except Exception:
                continue

    def _add_page_numbers(self, section, position: Optional[PageNumberPosition]):
        """Add page numbers to document"""
        if position is None:
            position = PageNumberPosition.BOTTOM_CENTER

        is_header = position.value.startswith("top")

        if is_header:
            target = section.header
        else:
            target = section.footer

        target.is_linked_to_previous = False

        if target.paragraphs:
            paragraph = target.paragraphs[0]
        else:
            paragraph = target.add_paragraph()

        if "left" in position.value:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif "center" in position.value:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "right" in position.value:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def get_document_preview(self, file_id: str) -> dict:
        """Get a text preview of the document"""
        file_path = self.get_file_path(file_id)
        if not file_path:
            formatted_path = os.path.join(self.upload_dir, f"{file_id}_formatted.docx")
            if os.path.exists(formatted_path):
                file_path = formatted_path
            else:
                raise FileNotFoundError(f"File not found: {file_id}")

        doc = Document(file_path)

        content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)

        return {
            "content": "\n\n".join(content[:50]),
            "page_count": len(doc.sections),
            "paragraph_count": len(doc.paragraphs),
        }

    def delete_file(self, file_id: str):
        """Delete a file by its ID"""
        file_path = self.get_file_path(file_id)
        if file_path and os.path.exists(file_path):
            os.remove(file_path)

        formatted_path = os.path.join(self.upload_dir, f"{file_id}_formatted.docx")
        if os.path.exists(formatted_path):
            os.remove(formatted_path)


# Singleton instance
document_processor = DocumentProcessor()
