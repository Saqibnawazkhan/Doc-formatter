import sys
import io
from pathlib import Path
from docx import Document

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Add current directory to path for imports
sys.path.insert(0, str(Path(__file__).parent))

from services.document_processor import DocumentProcessor
from models.formatting_options import (
    FormattingOptions,
    TextFormattingOptions,
    DocumentStructureOptions,
    FontFamily,
)

def test_document_formatting():
    """Test document formatting with the user's file"""

    # Initialize processor
    processor = DocumentProcessor(upload_dir="test_uploads")

    # Read the test file
    input_file = r"c:\Users\saqib\Downloads\formatted_document (6).docx"

    print("=" * 80)
    print(f"TESTING DOCUMENT FORMATTING")
    print("=" * 80)
    print(f"Input file: {input_file}")

    # Read the original file content
    with open(input_file, 'rb') as f:
        file_content = f.read()

    # Save it to processor
    file_id = processor.save_uploaded_file(file_content, "formatted_document (6).docx")
    print(f"Uploaded file ID: {file_id}")

    # Create formatting options - test with heading size 14 as user mentioned
    options = FormattingOptions(
        text=TextFormattingOptions(
            font_family=FontFamily.CALIBRI,
            font_size=12,
        ),
        structure=DocumentStructureOptions(
            normalize_headings=True,
            create_toc=True,
            heading_font_family=FontFamily.CALIBRI,
            h1_size=18,
            h1_color="#000000",
            h1_bold=True,
            h2_size=16,
            h2_color="#000000",
            h2_bold=True,
            h3_size=14,  # User mentioned 14
            h3_color="#000000",
            h3_bold=True,
        )
    )

    print("\nFormatting options:")
    print(f"  Text: {options.text.font_family.value if options.text else 'None'}, {options.text.font_size if options.text else 'None'}pt")
    print(f"  H1: {options.structure.h1_size}pt, Bold: {options.structure.h1_bold}")
    print(f"  H2: {options.structure.h2_size}pt, Bold: {options.structure.h2_bold}")
    print(f"  H3: {options.structure.h3_size}pt, Bold: {options.structure.h3_bold}")
    print(f"  Create TOC: {options.structure.create_toc}")

    # Format the document
    print("\nFormatting document...")
    formatted_file_id = processor.format_document(file_id, options)
    print(f"Formatted file ID: {formatted_file_id}")

    # Load the formatted document
    formatted_path = processor.get_file_path(formatted_file_id)
    if not formatted_path:
        formatted_path = f"test_uploads/{formatted_file_id}_formatted.docx"

    print(f"\nAnalyzing formatted document: {formatted_path}")

    doc = Document(formatted_path)

    print("\n" + "=" * 80)
    print("FORMATTED DOCUMENT ANALYSIS")
    print("=" * 80)

    # Check for markdown characters
    markdown_chars_found = False
    print("\n1. Checking for markdown characters (**, ###, ---)...")
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text
        if '**' in text or '###' in text or text.strip() == '---':
            print(f"   ⚠️  Para {i}: Still has markdown: {text[:80]}")
            markdown_chars_found = True

    if not markdown_chars_found:
        print("   ✓ No markdown characters found in first 30 paragraphs!")

    # Check headings and their sizes
    print("\n2. Checking heading detection and sizes...")
    heading_count = 0
    for i, para in enumerate(doc.paragraphs):
        if para.style and 'heading' in para.style.name.lower():
            heading_count += 1

            # Get font size from first run
            if para.runs:
                font_size = para.runs[0].font.size.pt if para.runs[0].font.size else "None"
                is_bold = para.runs[0].font.bold

                print(f"   Para {i}: {para.style.name} - {font_size}pt - Bold: {is_bold} - {para.text[:60]}")
            else:
                print(f"   Para {i}: {para.style.name} - No runs - {para.text[:60]}")

    print(f"\n   Total headings: {heading_count}")

    # Check for Table of Contents
    print("\n3. Checking for Table of Contents...")
    if doc.paragraphs and doc.paragraphs[0].text.strip() == "Table of Contents":
        print("   ✓ Table of Contents found at the beginning!")
    else:
        print("   ⚠️  No Table of Contents found")

    # Save to desktop for manual inspection
    output_file = r"c:\Users\saqib\Desktop\test_formatted_output.docx"
    try:
        doc.save(output_file)
        print(f"\n✓ Formatted document saved to: {output_file}")
    except Exception as e:
        print(f"\n⚠️  Could not save to desktop: {e}")

    print("\n" + "=" * 80)
    print("TEST COMPLETE")
    print("=" * 80)

if __name__ == "__main__":
    test_document_formatting()
