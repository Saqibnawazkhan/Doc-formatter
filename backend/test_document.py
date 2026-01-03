import sys
import io
from docx import Document
from pathlib import Path

# Force UTF-8 output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def analyze_document(file_path):
    """Analyze a Word document to understand its structure and content"""
    try:
        doc = Document(file_path)

        print(f"=" * 80)
        print(f"DOCUMENT ANALYSIS: {Path(file_path).name}")
        print(f"=" * 80)
        print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
        print(f"Total sections: {len(doc.sections)}")

        # Analyze paragraphs
        print(f"\n{'=' * 80}")
        print("PARAGRAPH ANALYSIS (First 30 paragraphs)")
        print(f"{'=' * 80}\n")

        for i, para in enumerate(doc.paragraphs[:30]):
            text = para.text.strip()
            if not text:
                print(f"Para {i}: [EMPTY]")
                continue

            # Get style info
            style_name = para.style.name if para.style else "No Style"

            # Get formatting info from first run
            if para.runs:
                first_run = para.runs[0]
                font_size = first_run.font.size.pt if first_run.font.size else "None"
                font_name = first_run.font.name if first_run.font.name else "None"
                is_bold = first_run.font.bold
                is_italic = first_run.font.italic

                print(f"\nPara {i}:")
                print(f"  Style: {style_name}")
                print(f"  Font: {font_name}, Size: {font_size}pt, Bold: {is_bold}, Italic: {is_italic}")
                print(f"  Text (first 100 chars): {text[:100]}")

                # Check for special characters
                special_chars = [c for c in text if c in '**###...']
                if special_chars:
                    print(f"  ⚠️  Special characters found: {special_chars}")
            else:
                print(f"Para {i}: [No runs] {text[:100]}")

        # Check for special character patterns
        print(f"\n{'=' * 80}")
        print("SPECIAL CHARACTER PATTERNS")
        print(f"{'=' * 80}\n")

        patterns_found = {
            '**': [],
            '###': [],
            '...': [],
        }

        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if '**' in text:
                patterns_found['**'].append((i, text[:100]))
            if '###' in text:
                patterns_found['###'].append((i, text[:100]))
            if '...' in text:
                patterns_found['...'].append((i, text[:100]))

        for pattern, occurrences in patterns_found.items():
            if occurrences:
                print(f"\nPattern '{pattern}' found in {len(occurrences)} paragraphs:")
                for para_idx, text in occurrences[:5]:  # Show first 5
                    print(f"  Para {para_idx}: {text}")

        # Check heading detection
        print(f"\n{'=' * 80}")
        print("HEADING DETECTION")
        print(f"{'=' * 80}\n")

        heading_count = 0
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""

            # Check if it's a heading style
            if 'heading' in style_name.lower():
                heading_count += 1
                print(f"Para {i}: {style_name} - {para.text[:80]}")

            # Check if it looks like a heading (bold, short, larger font)
            elif para.runs:
                first_run = para.runs[0]
                is_bold = first_run.font.bold
                font_size = first_run.font.size.pt if first_run.font.size else 12
                text_length = len(para.text.strip())

                if is_bold and text_length < 100 and text_length > 0:
                    print(f"Para {i}: Potential heading (Bold, {font_size}pt, {text_length} chars) - {para.text[:80]}")

        print(f"\nTotal headings found: {heading_count}")

    except Exception as e:
        print(f"Error analyzing document: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = r"c:\Users\saqib\Downloads\formatted_document (6).docx"

    analyze_document(file_path)
